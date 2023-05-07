#PDF documents
import PyPDF2
#Word documents
import docx
#For folders
import os
#For graphic interface
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text

# Feedback document
def create_template_feedback():
    doc_f = docx.Document()
    title_f = doc_f.add_paragraph("Feedback ERGOB1070 Technology and Society")
    title_f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_f.style = "Title"
    doc_f.add_paragraph()
    return doc_f

def add_feedback(text, doc_feedback):
    doc_feedback.add_paragraph(text)
    directory_feedbacks = get_folder_path()
    doc_feedback.save(os.path.join(directory_feedbacks, "Feedback.docx"))

def find_student_number(text):
    words = text.split()
    word_index = 0
    number_found = False
    student_number = None
    while not number_found and word_index<len(words)-2:
        if words[word_index] == "Candidate" and words[word_index+1] == "number:" and words[word_index+2].isdigit():
            number_found = True
            student_number = words[word_index+2]
            while words[word_index+3].isdigit():
                student_number = student_number + words[word_index+3]
                word_index += 1
        word_index += 1
    return number_found, student_number

def identify_answers(text):
    answers = []
    start_identifier, end_identifier = get_identifiers()
    start_index = 0
    while True:
        start_index = text.find(start_identifier, start_index)
        if start_index == -1:
            break
        end_index = text.find(end_identifier, start_index + len(start_identifier))
        if end_index == -1:
            break
        answer = text[start_index + len(start_identifier):end_index].strip()
        answer = ''.join(filter(lambda x: x.isprintable(), answer))
        answers.append(answer)
        start_index = end_index + len(end_identifier)
    return answers

def check_identifiers(text_exam, start_identifier, end_identifier):
    nb_start_identifier = text_exam.count(start_identifier)
    nb_end_identifier = text_exam.count(end_identifier)
    if start_identifier == end_identifier:
        identifier_ok = (nb_start_identifier%2 == 0)
    else:
        identifier_ok = (nb_start_identifier == nb_end_identifier)
    return identifier_ok

def check_answer(text_answer, accuracy, wordcount):
    nb_words = len(text_answer.split())
    if accuracy != None:
        min_words = int(wordcount - accuracy*wordcount)
        max_words = int(wordcount + accuracy * wordcount)
    if nb_words == 0:
        question_failed = True
        text_feedback = "No answer for this question"
    elif accuracy != None and not(min_words < nb_words < max_words):
        question_failed = True
        text_feedback = "You didn't respect the number of words expected. You wrote " \
                        + str(nb_words) + " word(s). You were supposed to write between " \
                        + str(min_words) + " and " + str(max_words) + " words"
    else:
        question_failed = False
        text_feedback = None

    return question_failed, text_feedback

def correct_exam(answers, doc_feedback_exam):
    word_accuracy = return_accuracy_wordcount()
    wordcount = return_wordcount()
    is_F = False
    for i in range(len(answers)):
        is_F, feedback = is_F or check_answer(answers[i], word_accuracy, wordcount[i])[0], \
            check_answer(answers[i], word_accuracy, wordcount[i])[1]
        if feedback != None:
            add_feedback("Question " + str(i+1) + ": " + feedback, doc_feedback_exam)
    if not is_F:
        add_feedback("OK", doc_feedback_exam)

def split_exams(text):
    exams = []
    start_tag, end_tag = get_tags_exam()
    start_index = 0
    while True:
        start_index = text.find(start_tag, start_index)
        if start_index == -1:
            break
        end_index = text.find(end_tag, start_index + len(start_tag))
        if end_index == -1:
            break
        exam = text[start_index + len(start_tag):end_index].strip()
        exam = ''.join(filter(lambda x: x.isprintable(), exam))
        exams.append(exam)
        start_index = end_index + len(end_tag)
    return exams

def correct_all_exams():
    whole_text = extract_text_from_pdf(get_folder_path() + "/Exams.pdf")
    exams_to_correct = split_exams(whole_text)
    doc_feedback = create_template_feedback()
    text_one_exam = exams_to_correct[0]
    alert_user(whole_text, text_one_exam)
    wordcount_window(nb_questions(text_one_exam)[0])
    for exam in exams_to_correct:
        is_found, number = find_student_number(exam)
        if is_found:
            add_feedback("Student number: " + number, doc_feedback)
        else:
            add_feedback("Student number: Not found", doc_feedback)

        answers_from_text = (identify_answers(exam))
        correct_exam(answers_from_text, doc_feedback)

#GUI
def user_window():

    def get_input():
        global input_start_identifier
        global input_end_identifier
        global input_start_tag
        global input_end_tag
        #Get values of entries
        input_start_identifier = entry_start_identifier.get()
        input_end_identifier = entry_end_identifier.get()
        input_start_tag = entry_start_tag.get()
        input_end_tag = entry_end_tag.get()
        window.destroy()

    def select_folder():
        folder_path = filedialog.askdirectory()
        return folder_path

    def on_select_folder():
        global folder_path
        folder_path = select_folder()
        folder_path_label.config(text="Selected folder : " + folder_path)

    # Create the window
    window = tk.Tk()
    window.title("Automated Exam Assessment")
    window.geometry("500x300")  # set window size

    #Folder
    folder_button = tk.Button(window, text="Select Folder", command=on_select_folder)
    folder_button.pack()
    folder_path_label = tk.Label(window, text="")
    folder_path_label.pack()

    #Entry to get identifiers
    #Start identifier
    start_identifier_label = tk.Label(window, text="Start identifier")
    start_identifier_label.pack(pady=5)
    entry_start_identifier = tk.Entry(window)
    entry_start_identifier.pack()
    #End identifier
    end_identifier_label = tk.Label(window, text="End identifier")
    end_identifier_label.pack(pady=5)
    entry_end_identifier = tk.Entry(window)
    entry_end_identifier.pack()
    #Start tag exam
    start_tag_label = tk.Label(window, text="Start tag exam")
    start_tag_label.pack(pady=5)
    entry_start_tag = tk.Entry(window)
    entry_start_tag.pack()
    #End tag exam
    end_tag_label = tk.Label(window, text="End tag exam")
    end_tag_label.pack(pady=5)
    entry_end_tag = tk.Entry(window)
    entry_end_tag.pack()

    OK_button = tk.Button(window, text="OK",command=get_input)
    OK_button.pack(pady=20)

    window.mainloop()

def get_identifiers():
    return input_start_identifier, input_end_identifier

def get_tags_exam():
    return input_start_tag, input_end_tag

def get_folder_path():
    return folder_path


#Infos & Alerts for user
def nb_questions(text):
    start_identifier, end_identifier = get_identifiers()
    nb_start_identifiers = text.count(start_identifier)
    nb_end_identifiers = text.count(end_identifier)
    return nb_start_identifiers, nb_end_identifiers

def nb_exams(text):
    start_tag, end_tag = get_tags_exam()
    nb_start_tag = text.count(start_tag)
    nb_end_tag = text.count(end_tag)
    return nb_start_tag, nb_end_tag

def alert_user(full_text, exam_text):
    start_q, end_q = nb_questions(exam_text)
    start_t, end_t = nb_exams(full_text)
    if start_q == end_q:
        messagebox.showinfo("Info", str(start_q) + " questions have been detected per exam")
    else:
        messagebox.showwarning("Warning", str(start_q) + " start identifiers and " + str(end_q)
                               + " end identifiers have been detected. Thus questions can't be identified clearly.")
    if start_t == end_t:
        messagebox.showinfo("Info", str(start_t) + " exams have been detected")
    else:
        messagebox.showwarning("Warning", str(start_t) + " start tags and " + str(end_t)
                               + " end tags have been detected. Thus exams can't be identified clearly.")

def wordcount_window(nb_questions):

    def get_values():
        global wordcounts
        global accuracy_wc
        wordcounts = []
        for entry in entries:
            value = entry.get()
            wordcounts.append(value)
        accuracy_wc = entry_accuracy.get()
        window_wc.destroy()

    # Create the window
    window_wc = tk.Tk()
    window_wc.title("Wordcount for each question")
    window_wc.geometry("500x300")  # set window size

    entries_frame = tk.Frame(window_wc)
    entries_frame.pack(side="left", fill="y")

    scrollbar = tk.Scrollbar(entries_frame, orient="vertical")
    scrollbar.pack(side="right", fill="y")

    canvas = tk.Canvas(entries_frame, yscrollcommand=scrollbar.set)
    canvas.pack(side="left", fill="both", expand=True)

    scrollbar.config(command=canvas.yview)

    inner_frame = tk.Frame(canvas)
    inner_frame_id = canvas.create_window((0, 0), window=inner_frame, anchor="nw")

    label = tk.Label(inner_frame, text="Accuracy (%)")
    label.pack(pady=5)

    entry_accuracy = tk.Entry(inner_frame)
    entry_accuracy.pack()

    entries = []
    for i in range(nb_questions):
        label = tk.Label(inner_frame, text=f"Question {i + 1}")
        label.pack()

        entry = tk.Entry(inner_frame)
        entry.pack()

        entries.append(entry)

    button = tk.Button(window_wc, text="Get Values", command=get_values)
    button.pack()

    def resize_canvas(event):
        canvas.config(scrollregion=canvas.bbox("all"))

    canvas.bind("<Configure>", resize_canvas)

    def resize_inner_frame(event):
        canvas.itemconfig(inner_frame_id, width=event.width)

    inner_frame.bind("<Configure>", resize_inner_frame)

    window_wc.mainloop()

def return_wordcount():
    wordcounts_converted = wordcounts
    for j in range(len(wordcounts_converted)):
        if wordcounts_converted[j] != None:
            try:
                wordcounts_converted[j] = float(wordcounts_converted[j])
            except ValueError:
                wordcounts_converted[j] = None
    return wordcounts_converted

def return_accuracy_wordcount():
     accuracy = accuracy_wc
     if accuracy != None:
         try:
             accuracy = float(accuracy)
         except ValueError:
             accuracy = None
     return accuracy/100 #Conversion percentage

user_window()
correct_all_exams()
#print(extract_text_from_pdf("Exams.pdf"))